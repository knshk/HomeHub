"""Indexing + semantic search.

- Extract text from txt/md/pdf/docx.
- Chunk text; embed each chunk via Ollama.
- For photos: caption via vision model, then embed the caption.
- Store chunks with float32 embedding BLOBs.
- Cosine search with numpy over file_chunks, scoped by authz, plus keyword
  fallback. Returns ranked [{file_id,filename,kind,score,snippet}].
"""
import struct
from pathlib import Path

import numpy as np

from . import config, db, integration
from .errors import HubError


# ----------------------------------------------------------------------------
# Embedding (de)serialization: float32 bytes
# ----------------------------------------------------------------------------
def embedding_to_blob(vec: list[float]) -> bytes:
    arr = np.asarray(vec, dtype=np.float32)
    return arr.tobytes()


def blob_to_embedding(blob: bytes) -> np.ndarray:
    if not blob:
        return np.zeros((0,), dtype=np.float32)
    return np.frombuffer(blob, dtype=np.float32)


# ----------------------------------------------------------------------------
# Text extraction
# ----------------------------------------------------------------------------
def _ext(filename: str) -> str:
    return Path(filename).suffix.lower().lstrip(".")


def extract_text(path: str, filename: str, mime: str) -> str:
    ext = _ext(filename)
    try:
        if ext in ("txt", "md", "markdown", "csv", "log", "json", "text"):
            return Path(path).read_text(encoding="utf-8", errors="replace")
        if ext == "pdf" or mime == "application/pdf":
            return _extract_pdf(path)
        if ext in ("docx",) or mime in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ):
            return _extract_docx(path)
        # Unknown text-ish type: try utf-8 decode, else empty.
        try:
            return Path(path).read_text(encoding="utf-8", errors="strict")
        except Exception:
            return ""
    except Exception:
        return ""


def _extract_pdf(path: str) -> str:
    from pypdf import PdfReader
    reader = PdfReader(path)
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n".join(parts)


def _extract_docx(path: str) -> str:
    import docx  # python-docx
    document = docx.Document(path)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


# ----------------------------------------------------------------------------
# Chunking
# ----------------------------------------------------------------------------
def chunk_text(text: str, size: int | None = None, overlap: int | None = None) -> list[str]:
    text = (text or "").strip()
    if not text:
        return []
    size = size or config.CHUNK_SIZE
    overlap = overlap if overlap is not None else config.CHUNK_OVERLAP
    if overlap >= size:
        overlap = size // 4
    chunks = []
    start = 0
    n = len(text)
    while start < n:
        end = min(start + size, n)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= n:
            break
        start = end - overlap
    return chunks


PHOTO_EXTS = {"jpg", "jpeg", "png", "gif", "webp", "bmp", "heic", "heif", "tif", "tiff"}


def is_photo(filename: str, mime: str) -> bool:
    if mime and mime.startswith("image/"):
        return True
    return _ext(filename) in PHOTO_EXTS


# ----------------------------------------------------------------------------
# Indexing
# ----------------------------------------------------------------------------
async def index_file(file_id: int) -> None:
    """Index a single file row: extract/caption -> chunk -> embed -> store.

    Best-effort: failures mark indexed=0 but do not raise to the request path
    unless explicitly desired. Here we swallow upstream errors so an upload still
    succeeds even if Ollama is down (search just won't find it yet).
    """
    conn = db.connect()
    try:
        row = db.query_one(conn, "SELECT * FROM files WHERE id=?", (file_id,))
        if row is None:
            return
        rec = dict(row)
        kind = rec["kind"]
        path = rec["stored_path"]
        filename = rec["filename"]
        mime = rec["mime"]

        # Clear any previous chunks (re-index safe).
        db.execute(conn, "DELETE FROM file_chunks WHERE file_id=?", (file_id,))
        conn.commit()

        caption_text = ""
        chunks: list[str] = []

        if kind == "photo":
            try:
                with open(path, "rb") as f:
                    image_bytes = f.read()
                caption_text = await integration.caption(image_bytes)
            except Exception:
                caption_text = ""
            if caption_text:
                db.execute(conn, "UPDATE files SET caption=? WHERE id=?", (caption_text, file_id))
                conn.commit()
                chunks = [caption_text]
        else:
            text = extract_text(path, filename, mime)
            chunks = chunk_text(text)

        stored = 0
        for idx, ch in enumerate(chunks):
            try:
                vec = await integration.embeddings(ch)
                blob = embedding_to_blob(vec)
            except Exception:
                blob = None
            db.execute(
                conn,
                "INSERT INTO file_chunks (file_id, chunk_index, text, embedding) VALUES (?,?,?,?)",
                (file_id, idx, ch, blob),
            )
            stored += 1
        conn.commit()

        indexed_flag = 1 if stored > 0 else 0
        db.execute(conn, "UPDATE files SET indexed=? WHERE id=?", (indexed_flag, file_id))
        conn.commit()
    finally:
        conn.close()


# ----------------------------------------------------------------------------
# Search
# ----------------------------------------------------------------------------
def _cosine(query: np.ndarray, mat: np.ndarray) -> np.ndarray:
    if mat.size == 0:
        return np.zeros((0,), dtype=np.float32)
    qn = np.linalg.norm(query)
    if qn == 0:
        return np.zeros((mat.shape[0],), dtype=np.float32)
    mn = np.linalg.norm(mat, axis=1)
    mn[mn == 0] = 1e-9
    return (mat @ query) / (mn * qn)


def _snippet(text: str, query: str, length: int = 200) -> str:
    text = text or ""
    if not text:
        return ""
    low = text.lower()
    terms = [t for t in query.lower().split() if t]
    pos = -1
    for t in terms:
        pos = low.find(t)
        if pos != -1:
            break
    if pos == -1:
        pos = 0
    start = max(0, pos - 40)
    snip = text[start:start + length].strip().replace("\n", " ")
    if start > 0:
        snip = "…" + snip
    if start + length < len(text):
        snip = snip + "…"
    return snip


def _accessible_file_ids(conn, username: str, role: str, kind: str | None) -> dict[int, dict]:
    """Return {file_id: file_row_dict} the user may read: owned OR shared.
    Optionally filtered by kind. Admins see everything."""
    params: list = []
    where = []
    if role == "admin":
        pass
    else:
        where.append("(owner_username=? OR shared=1)")
        params.append(username)
    if kind in ("file", "photo"):
        where.append("kind=?")
        params.append(kind)
    sql = "SELECT * FROM files"
    if where:
        sql += " WHERE " + " AND ".join(where)
    rows = db.query_all(conn, sql, params)
    return {r["id"]: dict(r) for r in rows}


async def search(username: str, role: str, q: str, kind: str | None = None,
                 top_k: int = 20) -> list[dict]:
    """Semantic (embeddings) + keyword search scoped to accessible files."""
    q = (q or "").strip()
    if not q:
        return []

    conn = db.connect()
    try:
        accessible = _accessible_file_ids(conn, username, role, kind)
        if not accessible:
            return []
        file_ids = list(accessible.keys())

        # Pull chunks for accessible files only.
        placeholders = ",".join("?" for _ in file_ids)
        chunk_rows = db.query_all(
            conn,
            f"SELECT id, file_id, chunk_index, text, embedding FROM file_chunks "
            f"WHERE file_id IN ({placeholders})",
            file_ids,
        )
    finally:
        conn.close()

    # Try semantic query embedding (may fail if Ollama down -> keyword only).
    query_vec = None
    try:
        query_vec = np.asarray(await integration.embeddings(q), dtype=np.float32)
    except Exception:
        query_vec = None

    qterms = [t for t in q.lower().split() if t]

    # Score per chunk, keep best chunk per file.
    best: dict[int, dict] = {}
    sem_vecs = []
    sem_meta = []
    for cr in chunk_rows:
        fid = cr["file_id"]
        text = cr["text"] or ""
        # keyword score
        low = text.lower()
        kw = sum(low.count(t) for t in qterms)
        kw_score = 0.0
        if kw > 0:
            kw_score = min(1.0, 0.2 + 0.1 * kw)

        emb = blob_to_embedding(cr["embedding"])
        if query_vec is not None and emb.size == query_vec.size and emb.size > 0:
            sem_vecs.append(emb)
            sem_meta.append((fid, text, kw_score))
        else:
            # keyword-only candidate
            _consider(best, fid, kw_score, text, q)

    if sem_vecs:
        mat = np.vstack(sem_vecs)
        sims = _cosine(query_vec, mat)
        for (fid, text, kw_score), sim in zip(sem_meta, sims):
            score = 0.75 * float(sim) + 0.25 * kw_score
            _consider(best, fid, score, text, q)

    # Build ranked results.
    results = []
    for fid, info in best.items():
        frow = accessible.get(fid)
        if not frow:
            continue
        results.append({
            "file_id": fid,
            "filename": frow["filename"],
            "kind": frow["kind"],
            "score": round(float(info["score"]), 4),
            "snippet": _snippet(info["text"], q),
        })
    results.sort(key=lambda r: r["score"], reverse=True)
    return results[:top_k]


def _consider(best: dict, fid: int, score: float, text: str, q: str) -> None:
    cur = best.get(fid)
    if cur is None or score > cur["score"]:
        best[fid] = {"score": score, "text": text}
